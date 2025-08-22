"""
CV Processing Utilities
Shared functions for CV content extraction and review processing
"""

import os
import json
import io
import base64
from typing import Optional
from googleapiclient.discovery import build
from google.oauth2 import service_account
from pdfminer.high_level import extract_text
from github import Github
import re
from backend.crew import CrewAI


def setup_drive_service():
    """Setup Google Drive service with credentials."""
    creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")
    if not creds_json:
        raise Exception("Google Drive credentials not configured")
    
    if creds_json.strip().startswith('{'):
        creds_dict = json.loads(creds_json)
    else:
        with open(creds_json) as f:
            creds_dict = json.load(f)
            
    creds = service_account.Credentials.from_service_account_info(
        creds_dict, 
        scopes=["https://www.googleapis.com/auth/drive", "https://www.googleapis.com/auth/drive.file"]
    )
    return build('drive', 'v3', credentials=creds)


def get_gdrive_cv_content(file_id):
    """Get CV content from Google Drive file."""
    creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")
    if creds_json.strip().startswith('{'):
        creds_dict = json.loads(creds_json)
    else:
        with open(creds_json) as f:
            creds_dict = json.load(f)
    creds = service_account.Credentials.from_service_account_info(creds_dict, scopes=["https://www.googleapis.com/auth/drive"])
    service = build('drive', 'v3', credentials=creds)
    file = service.files().get(fileId=file_id, fields="name, mimeType").execute()
    mime = file['mimeType']
    if mime == 'application/pdf':
        data = service.files().get_media(fileId=file_id).execute()
        text = extract_text(io.BytesIO(data))
        return text
    elif mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        data = service.files().get_media(fileId=file_id).execute()
        from docx import Document
        doc = Document(io.BytesIO(data))
        return '\n'.join([p.text for p in doc.paragraphs])
    return "Unsupported file type"


def run_openai_review(cv_text, cv_name, job_description=None):
    """Run OpenAI review using CrewAI."""
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        return "OpenAI API key not set."
    # Setup CrewAI agent with OpenAI
    inputs = {
        'cv_name': cv_name,
        'cv_text': cv_text,
        'job_description': job_description,
    }
    result = CrewAI().crew().kickoff(inputs=inputs)
    return result


def extract_fit_score(review_text: str) -> int:
    """
    Extract fit score from review text with multiple fallback patterns.
    Returns an integer between 0-100, defaulting to 0 if not found.
    """
    if not review_text:
        return 0
    
    # Multiple regex patterns to catch different formats
    patterns = [
        # "Fit Score: 85" or "fit score: 85" (case insensitive, at start of line)
        r"(?i)^fit\s*score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "Fit Score: 85" anywhere in text
        r"(?i)fit\s*score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "Score: 85" or "Overall Score: 85"
        r"(?i)(?:overall\s+)?score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "85/100" or "85 out of 100"
        r"(\d+(?:\.\d+)?)\s*(?:/100|out\s+of\s+100)",
        
        # Numbers followed by % (assuming it's out of 100)
        r"(\d+(?:\.\d+)?)%",
        
        # Just look for any number in parentheses like "(85)"
        r"\((\d+(?:\.\d+)?)\)",
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, review_text, re.MULTILINE)
        if matches:
            try:
                # Take the first match
                score = float(matches[0])
                
                # Ensure score is within reasonable bounds (0-100)
                if score > 100:
                    score = min(score, 100)  # Cap at 100
                elif score < 0:
                    score = 0
                
                return int(round(score))
            except (ValueError, TypeError) as e:
                print(f"[WARN] Failed to parse score '{matches[0]}': {e}")
                continue
    
    # If no patterns match, try to find any number that might be a score
    # Look for standalone numbers between 0-100
    standalone_numbers = re.findall(r'\b(\d+(?:\.\d+)?)\b', review_text)
    for num_str in standalone_numbers:
        try:
            num = float(num_str)
            if 0 <= num <= 100:
                print(f"[INFO] Using standalone number {num} as potential fit score")
                return int(round(num))
        except ValueError:
            continue
    
    print(f"[WARN] No fit score found in review text: {review_text[:200]}...")
    return 0


def move_to_qualified_folder(file_id: str, qualified_folder_id: str) -> bool:
    """Move the file to the qualified candidates folder in Google Drive."""
    try:
        creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")
        
        if not creds_json or not qualified_folder_id:
            print("[ERROR] Google Drive credentials or qualified folder ID not set")
            return False
            
        if creds_json.strip().startswith('{'):
            creds_dict = json.loads(creds_json)
        else:
            with open(creds_json) as f:
                creds_dict = json.load(f)
                
        creds = service_account.Credentials.from_service_account_info(
            creds_dict, 
            scopes=["https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]
        )
        service = build('drive', 'v3', credentials=creds)
        
        try:
            # Get current file metadata including parents
            file = service.files().get(
                fileId=file_id,
                fields='parents,name'
            ).execute()
            
            # Remove the file from its current folder
            previous_parents = ",".join(file.get('parents', []))
            
            # Move the file to the qualified folder
            updated_file = service.files().update(
                fileId=file_id,
                addParents=qualified_folder_id,
                removeParents=previous_parents,
                fields='id, parents, name'
            ).execute()
            
            print(f"[INFO] Successfully moved file {file.get('name')} to qualified folder")
            return True
            
        except Exception as e:
            print(f"[ERROR] Error accessing file {file_id}: {str(e)}")
            return False
        
        return True
    except Exception as e:
        print(f"[ERROR] Failed to move file {file_id} to qualified folder: {str(e)}")
        return False


def get_file_content(service, file_id):
    """Helper function to get the content of a file from Google Drive."""
    try:
        # Get file metadata to check its type
        file = service.files().get(fileId=file_id, fields="mimeType, name").execute()
        mime_type = file.get('mimeType', '')
        
        # Get the file content
        data = service.files().get_media(fileId=file_id).execute()
        
        if mime_type == 'application/pdf':
            # Handle PDF files
            return extract_text(io.BytesIO(data))
            
        elif mime_type == 'application/vnd.google-apps.document':
            # For Google Docs, use the export method
            data = service.files().export(
                fileId=file_id,
                mimeType='text/plain'
            ).execute()
            return data.decode('utf-8') if isinstance(data, bytes) else str(data)
            
        else:
            # For regular text files
            if isinstance(data, bytes):
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return data.decode('utf-8', errors='ignore')
            return str(data)
            
    except Exception as e:
        raise Exception(f"Error reading file: {str(e)}")


def get_job_description_from_drive(folder_id=None):
    """Get job description from a file in Google Drive."""
    creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")
    if not creds_json:
        return "No Google Drive credentials configured"
    
    if creds_json.strip().startswith('{'):
        creds_dict = json.loads(creds_json)
    else:
        with open(creds_json) as f:
            creds_dict = json.load(f)
            
    creds = service_account.Credentials.from_service_account_info(
        creds_dict, 
        scopes=["https://www.googleapis.com/auth/drive"]
    )
    service = build('drive', 'v3', credentials=creds)

    try:
        if folder_id:
            results = service.files().list(
                q=f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') and trashed=false",
                fields="files(id, name, mimeType)",
                pageSize=1,  # We only need the first matching file
                orderBy="createdTime desc"  # Get the most recent file if multiple exist
            ).execute()

            files = results.get('files', [])
            if not files:
                return "No job description file found in the selected folder"

            # Use the first matching file
            file_id = files[0]['id']
            print(f"[INFO] Found job description file: {files[0]['name']}")
            return get_file_content(service, file_id)
        else:
            return 'No folder ID provided. No job description file found in the selected folder.'

    except Exception as e:
        error_msg = f"Error reading job description: {str(e)}"
        print(f"[ERROR] {error_msg}")
        return error_msg


def get_github_cv_content(path):
    """Get CV content from GitHub repository."""
    token = os.getenv("GITHUB_TK")
    repo_name = os.getenv("GITHUB_REPO")
    if not token or not repo_name:
        return ""
    g = Github(token)
    repo = g.get_repo(repo_name)
    file_content = repo.get_contents(path)
    if file_content.name.lower().endswith('.pdf'):
        data = base64.b64decode(file_content.content)
        text = extract_text(io.BytesIO(data))
        return text
    elif file_content.name.lower().endswith('.docx'):
        from docx import Document
        data = base64.b64decode(file_content.content)
        doc = Document(io.BytesIO(data))
        return '\n'.join([p.text for p in doc.paragraphs])
    return "Unsupported file type"


def run_gemini_review(cv_text, cv_name, job_description=None):
    """Run Gemini review using CrewAI."""
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        return "Gemini API key not set."
    # Setup CrewAI agent with Gemini
    inputs = {
        'cv_name': cv_name,
        'cv_text': cv_text,
        'job_description': job_description,
    }
    result = CrewAI().crew().kickoff(inputs=inputs)
    return result
