import os
from fastapi import FastAPI, Query, Body
from typing import List, Optional
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from googleapiclient.discovery import build
from google.oauth2 import service_account
from github import Github
from fastapi.responses import PlainTextResponse
import base64
from crewai import Crew, Agent, Task
import re
import json
from backend.crew import CrewAI
from backend.promptfoo_integration import PromptfooEvaluator, CVEvaluationAgent
import asyncio
from concurrent.futures import ThreadPoolExecutor
import io
from pdfminer.high_level import extract_text
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Google Drive configuration
GOOGLE_DRIVE_SCOPES = [
    "https://www.googleapis.com/auth/drive.readonly",
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/drive"
]

def get_google_drive_credentials():
    """
    Get Google Drive credentials from environment variable or file.
    Returns service account credentials object.
    
    Raises:
        ValueError: If credentials are not found or invalid
        FileNotFoundError: If credentials file is specified but not found
        json.JSONDecodeError: If credentials JSON is malformed
    """
    creds_json = os.getenv("GOOGLE_DRIVE_CREDENTIALS")
    
    if not creds_json:
        # Try to load from default file location
        default_creds_file = "google_drive_credentials.json"
        if os.path.exists(default_creds_file):
            logger.info(f"Loading Google Drive credentials from {default_creds_file}")
            with open(default_creds_file, 'r') as f:
                creds_dict = json.load(f)
        else:
            raise ValueError(
                "Google Drive credentials not found. Set GOOGLE_DRIVE_CREDENTIALS environment variable "
                "or place credentials in google_drive_credentials.json file."
            )
    else:
        # Parse credentials from environment variable
        if creds_json.strip().startswith('{'):
            # Direct JSON string
            try:
                creds_dict = json.loads(creds_json)
                logger.info("Loaded Google Drive credentials from environment variable (JSON string)")
            except json.JSONDecodeError as e:
                raise json.JSONDecodeError(f"Invalid JSON in GOOGLE_DRIVE_CREDENTIALS: {e}")
        else:
            # File path
            if not os.path.exists(creds_json):
                raise FileNotFoundError(f"Google Drive credentials file not found: {creds_json}")
            logger.info(f"Loading Google Drive credentials from file: {creds_json}")
            with open(creds_json, 'r') as f:
                creds_dict = json.load(f)
    
    # Validate required fields in service account credentials
    required_fields = ['type', 'project_id', 'private_key_id', 'private_key', 'client_email', 'client_id']
    missing_fields = [field for field in required_fields if field not in creds_dict]
    if missing_fields:
        raise ValueError(f"Missing required fields in Google Drive credentials: {missing_fields}")
    
    if creds_dict.get('type') != 'service_account':
        raise ValueError("Google Drive credentials must be for a service account")
    
    return service_account.Credentials.from_service_account_info(
        creds_dict, 
        scopes=GOOGLE_DRIVE_SCOPES
    )

def get_google_drive_service():
    """
    Get authenticated Google Drive service instance.
    
    Returns:
        googleapiclient.discovery.Resource: Authenticated Drive service
        
    Raises:
        ValueError: If credentials are invalid
        Exception: If service creation fails
    """
    try:
        credentials = get_google_drive_credentials()
        service = build('drive', 'v3', credentials=credentials)
        logger.info("Successfully created Google Drive service")
        return service
    except Exception as e:
        logger.error(f"Failed to create Google Drive service: {e}")
        raise

# Placeholder for CV metadata
class CVInfo(BaseModel):
    name: str
    source: str  # 'gdrive' or 'github'
    path: str
    job_description: Optional[str] = None

# Placeholder for review result
class ReviewResult(BaseModel):
    cv_name: str
    review: str

class ReviewAllResult(BaseModel):
    cv_name: str
    review: str
    fit_score: int
    token_count: Optional[int] = 0
    prompt_tokens: Optional[int] = 0
    completion_tokens: Optional[int] = 0
    total_tokens: Optional[int] = 0

def list_gdrive_folders(parent_id: str = None):
    """
    List folders in Google Drive.
    
    Args:
        parent_id: Parent folder ID. If None, uses GOOGLE_DRIVE_FOLDER_ID from environment
        
    Returns:
        List of folder dictionaries with 'id' and 'name' keys
    """
    try:
        root_folder_id = parent_id if parent_id is not None else os.getenv("GOOGLE_DRIVE_FOLDER_ID")
        if not root_folder_id:
            logger.warning("No Google Drive folder ID provided")
            return []
        
        service = get_google_drive_service()
        
        # Query for folders within the root folder
        results = service.files().list(
            q=f"'{root_folder_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false",
            fields="files(id, name)",
            pageSize=50
        ).execute()
        
        folders = results.get('files', [])
        logger.info(f"Found {len(folders)} folders in {root_folder_id}")
        return [{"id": folder["id"], "name": folder["name"]} for folder in folders]
        
    except Exception as e:
        logger.error(f"Error listing Google Drive folders: {e}")
        return []

def list_gdrive_cvs(folder_id=None):
    """
    List CV files (PDF and DOCX) in a Google Drive folder.
    
    Args:
        folder_id: Folder ID to search in. If None, uses GOOGLE_DRIVE_FOLDER_ID from environment
        
    Returns:
        List of CVInfo objects
    """
    try:
        if not folder_id:
            folder_id = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
        
        if not folder_id:
            logger.warning("No Google Drive folder ID provided for CV listing")
            return []
        
        service = get_google_drive_service()
        
        results = service.files().list(
            q=f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') and trashed=false",
            fields="files(id, name)",
            pageSize=50
        ).execute()
        
        files = results.get('files', [])
        logger.info(f"Found {len(files)} CV files in folder {folder_id}")
        return [CVInfo(name=f["name"], source="gdrive", path=f["id"]) for f in files]
        
    except Exception as e:
        logger.error(f"Error listing CV files from Google Drive: {e}")
        return []

def list_github_cvs():
    token = os.getenv("GITHUB_TK")
    repo_name = os.getenv("GITHUB_REPO")
    folder_path = os.getenv("GITHUB_FOLDER", "")
    if not token or not repo_name:
        return []
    g = Github(token)
    repo = g.get_repo(repo_name)
    files = repo.get_contents(folder_path)
    cv_files = []
    for f in files:
        if f.type == 'file' and (f.name.lower().endswith('.pdf') or f.name.lower().endswith('.docx')):
            cv_files.append(CVInfo(name=f.name, source="github", path=f.path))
    return cv_files

def move_to_qualified_folder(file_id: str, qualified_folder_id: str) -> bool:
    """Move the file to the qualified candidates folder in Google Drive."""
    if not qualified_folder_id:
        logger.error("No qualified folder ID provided")
        return False
        
    try:
        service = get_google_drive_service()
        
        # Get current file metadata including parents
        file = service.files().get(
            fileId=file_id,
            fields='parents,name'
        ).execute()
        
        # Remove the file from its current folder
        previous_parents = ",".join(file.get('parents', []))
        
        # Move the file to the qualified folder
        updated_file = service.files().update(
            fileId=file_id,
            addParents=qualified_folder_id,
            removeParents=previous_parents,
            fields='id, parents, name'
        ).execute()
        
        logger.info(f"Successfully moved file '{file.get('name')}' to qualified folder")
        return True
        
    except Exception as e:
        logger.error(f"Failed to move file {file_id} to qualified folder: {e}")
        return False

def move_to_unqualified_folder(file_id: str, unqualified_folder_id: str) -> bool:
    """Move the file to the unqualified candidates folder in Google Drive."""
    if not unqualified_folder_id:
        logger.error("No unqualified folder ID provided")
        return False
        
    try:
        service = get_google_drive_service()
        
        # Get current file metadata including parents
        file = service.files().get(
            fileId=file_id,
            fields='parents,name'
        ).execute()
        
        # Remove the file from its current folder
        previous_parents = ",".join(file.get('parents', []))
        
        # Move the file to the unqualified folder
        updated_file = service.files().update(
            fileId=file_id,
            addParents=unqualified_folder_id,
            removeParents=previous_parents,
            fields='id, parents, name'
        ).execute()
        
        logger.info(f"Successfully moved file '{file.get('name')}' to unqualified folder")
        return True
        
    except Exception as e:
        logger.error(f"Failed to move file {file_id} to unqualified folder: {e}")
        return False

def get_gdrive_cv_content(file_id):
    """
    Extract text content from a Google Drive file (PDF or DOCX).
    
    Args:
        file_id: Google Drive file ID
        
    Returns:
        str: Extracted text content
    """
    try:
        service = get_google_drive_service()
        
        # Get file metadata
        file = service.files().get(fileId=file_id, fields="name, mimeType").execute()
        mime_type = file['mimeType']
        file_name = file.get('name', 'Unknown')
        
        logger.info(f"Extracting content from file: {file_name} (type: {mime_type})")
        
        if mime_type == 'application/pdf':
            # Handle PDF files
            data = service.files().get_media(fileId=file_id).execute()
            text = extract_text(io.BytesIO(data))
            return text
            
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Handle DOCX files
            data = service.files().get_media(fileId=file_id).execute()
            from docx import Document
            doc = Document(io.BytesIO(data))
            return '\n'.join([p.text for p in doc.paragraphs])
            
        else:
            logger.warning(f"Unsupported file type: {mime_type} for file {file_name}")
            return f"Unsupported file type: {mime_type}"
            
    except Exception as e:
        logger.error(f"Error extracting content from Google Drive file {file_id}: {e}")
        return f"Error reading file: {str(e)}"

def get_github_cv_content(path):
    token = os.getenv("GITHUB_TK")
    repo_name = os.getenv("GITHUB_REPO")
    if not token or not repo_name:
        return ""
    g = Github(token)
    repo = g.get_repo(repo_name)
    file_content = repo.get_contents(path)
    if file_content.name.lower().endswith('.pdf'):
        import io
        from pdfminer.high_level import extract_text
        data = base64.b64decode(file_content.content)
        text = extract_text(io.BytesIO(data))
        return text
    elif file_content.name.lower().endswith('.docx'):
        import io
        from docx import Document
        data = base64.b64decode(file_content.content)
        doc = Document(io.BytesIO(data))
        return '\n'.join([p.text for p in doc.paragraphs])
    return "Unsupported file type"

def run_openai_review(cv_text, cv_name, job_description=None):
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        return "OpenAI API key not set."
    # Setup CrewAI agent with OpenAI
    inputs = {
        'cv_name': cv_name,
        'cv_text': cv_text ,
        'job_description': job_description,
    }
    result = CrewAI().crew().kickoff(inputs=inputs)
    return result

def run_gemini_review(cv_text, cv_name, job_description=None):
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        return "Gemini API key not set."
    # Setup CrewAI agent with Gemini
    inputs = {
        'cv_name': cv_name,
        'cv_text': cv_text ,
        'job_description': job_description,
    }
    result = CrewAI().crew().kickoff(inputs=inputs)
    return result

@app.get("/gdrive/folders")
def get_gdrive_folders(parent_id: Optional[str] = None):
    return list_gdrive_folders(parent_id=parent_id)

@app.get("/cvs", response_model=List[CVInfo])
def list_cvs(source: str = Query(..., regex="^(gdrive|github)$"), folder_id: Optional[str] = None):
    if source == 'gdrive':
        return list_gdrive_cvs(folder_id)
    else:
        return list_github_cvs()

@app.post("/review", response_model=ReviewResult)
def review_cv(cv: CVInfo):
    # Fetch CV content
    if cv.source == 'gdrive':
        cv_text = get_gdrive_cv_content(cv.path)
    elif cv.source == 'github':
        cv_text = get_github_cv_content(cv.path)
    else:
        return ReviewResult(cv_name=cv.name, review="Invalid source.")
    # Run CrewAI + OpenAI review, now with job description
    review = run_openai_review(cv_text, cv.name, cv.job_description)
    if hasattr(review, "raw"):
        review = review.raw
    return ReviewResult(cv_name=cv.name, review=review)

@app.post("/review_all", response_model=List[ReviewAllResult])
async def review_all_cvs(
    cvs: List[CVInfo],
    job_description: Optional[str] = Body(None),
    qualified_folder_id: Optional[str] = Body(None),
    unqualified_folder_id: Optional[str] = Body(None)
):

    # Create a single ThreadPoolExecutor for all tasks
    executor = ThreadPoolExecutor(max_workers=min(32, len(cvs) + 4))
    loop = asyncio.get_event_loop()
    
    async def process_cv(cv):
        try:
            # Step 2: Get CV text
            try:
                cv_text = await loop.run_in_executor(
                    executor,
                    get_gdrive_cv_content if cv.source == 'gdrive' else get_github_cv_content,
                    cv.path
                )
            except Exception as e:
                print(f"[ERROR] Failed to get CV text for {cv.name}: {str(e)}")
                return ReviewAllResult(
                    cv_name=cv.name,
                    review=f"Error getting CV content: {str(e)}",
                    fit_score=0,
                    token_count=0,
                    prompt_tokens=0,
                    completion_tokens=0,
                    total_tokens=0
                )

            # Step 3: Run OpenAI/CrewAI review
            try:
                review_output = await loop.run_in_executor(
                    executor,
                    run_openai_review,
                    cv_text,
                    cv.name,
                    job_description
                )
                review = review_output.raw if hasattr(review_output, "raw") else str(review_output)
                # Get token usage from review_output
                token_usage = getattr(review_output, "token_usage", None)
                if token_usage and hasattr(token_usage, "prompt_tokens"):
                    prompt_tokens = token_usage.prompt_tokens
                    completion_tokens = token_usage.completion_tokens
                    total_tokens = token_usage.total_tokens
                else:
                    # Fallback to estimating tokens if exact count not available
                    total_tokens = len(review.split()) * 2  # Rough estimate
                    prompt_tokens = len(cv_text.split())  # Rough estimate for input
                    completion_tokens = total_tokens - prompt_tokens
                token_count = total_tokens
            except Exception as e:
                print(f"[ERROR] Failed to review CV {cv.name}: {str(e)}")
                return ReviewAllResult(
                    cv_name=cv.name,
                    review=f"Error during CV review: {str(e)}",
                    fit_score=0,
                    token_count=0,
                    prompt_tokens=0,
                    completion_tokens=0,
                    total_tokens=0
                )

            # Step 4: Extract fit score with improved regex and parsing
            fit_score = extract_fit_score(review)

            # Debug logging
            print(f"[DEBUG] {cv.name} Fit Score: {fit_score}")
            
            # Step 5: Move qualified/unqualified CVs to respective folders
            moved_to_qualified = False
            moved_to_unqualified = False
            
            if cv.source == 'gdrive':
                if fit_score > 80 and qualified_folder_id:
                    # Move qualified CVs (score > 80) to qualified folder
                    try:
                        moved_to_qualified = await loop.run_in_executor(
                            executor,
                            move_to_qualified_folder,
                            cv.path,
                            qualified_folder_id
                        )
                        if moved_to_qualified:
                            print(f"[INFO] Moved qualified CV {cv.name} to qualified folder")
                    except Exception as e:
                        print(f"[ERROR] Failed to move CV {cv.name} to qualified folder: {str(e)}")
                        # Continue processing even if move fails
                        
                elif fit_score <= 80 and unqualified_folder_id:
                    # Move unqualified CVs (score <= 80) to unqualified folder
                    try:
                        moved_to_unqualified = await loop.run_in_executor(
                            executor,
                            move_to_unqualified_folder,
                            cv.path,
                            unqualified_folder_id
                        )
                        if moved_to_unqualified:
                            print(f"[INFO] Moved unqualified CV {cv.name} to unqualified folder")
                    except Exception as e:
                        print(f"[ERROR] Failed to move CV {cv.name} to unqualified folder: {str(e)}")
                        # Continue processing even if move fails

            # Prepare move status message
            move_status = ""
            if moved_to_qualified:
                move_status = "[Moved to qualified folder]"
            elif moved_to_unqualified:
                move_status = "[Moved to unqualified folder]"

            return ReviewAllResult(
                cv_name=cv.name,
                review=f"{review}\n\n{move_status}" if move_status else review,
                fit_score=fit_score,
                token_count=token_count,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                total_tokens=total_tokens
            )

        except Exception as e:
            # Don't let one error fail all
            print(f"[ERROR] Failed to process CV '{cv.name}': {e}")
            return ReviewAllResult(
                cv_name=cv.name,
                review=f"[Error processing CV: {str(e)}]",
                fit_score=0
            )
            
    if not cvs:
        return []

    try:
        # Process all CVs in parallel
        tasks = [process_cv(cv) for cv in cvs]
        try:
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Filter out any exceptions and convert them to error results
            processed_results = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    print(f"[ERROR] Task for CV {cvs[i].name} failed: {str(result)}")
                    processed_results.append(ReviewAllResult(
                        cv_name=cvs[i].name,
                        review=f"Error processing CV: {str(result)}",
                        fit_score=0
                    ))
                else:
                    processed_results.append(result)

            # Sort results by fit score (descending)
            sorted_results = sorted(processed_results, key=lambda x: x.fit_score, reverse=True)
            return sorted_results
        except Exception as e:
            print(f"[ERROR] Failed to gather results: {str(e)}")
            return [ReviewAllResult(
                cv_name="System Error",
                review=f"Failed to process CVs: {str(e)}",
                fit_score=0
            )]
    finally:
        # Ensure the executor is shut down properly
        try:
            executor.shutdown(wait=True)
        except Exception as e:
            print(f"[ERROR] Failed to shutdown executor: {str(e)}")


def extract_fit_score(review_text: str) -> int:
    """
    Extract fit score from review text with multiple fallback patterns.
    Returns an integer between 0-100, defaulting to 0 if not found.
    """
    if not review_text:
        return 0
    
    # Multiple regex patterns to catch different formats
    patterns = [
        # "Fit Score: 85" or "fit score: 85" (case insensitive, at start of line)
        r"(?i)^fit\s*score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "Fit Score: 85" anywhere in text
        r"(?i)fit\s*score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "Score: 85" or "Overall Score: 85"
        r"(?i)(?:overall\s+)?score\s*:\s*(\d+(?:\.\d+)?)",
        
        # "85/100" or "85 out of 100"
        r"(\d+(?:\.\d+)?)\s*(?:/100|out\s+of\s+100)",
        
        # Numbers followed by % (assuming it's out of 100)
        r"(\d+(?:\.\d+)?)%",
        
        # Just look for any number in parentheses like "(85)"
        r"\((\d+(?:\.\d+)?)\)",
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, review_text, re.MULTILINE)
        if matches:
            try:
                # Take the first match
                score = float(matches[0])
                
                # Ensure score is within reasonable bounds (0-100)
                if score > 100:
                    score = min(score, 100)  # Cap at 100
                elif score < 0:
                    score = 0
                
                return int(round(score))
            except (ValueError, TypeError) as e:
                print(f"[WARN] Failed to parse score '{matches[0]}': {e}")
                continue
    
    # If no patterns match, try to find any number that might be a score
    # Look for standalone numbers between 0-100
    standalone_numbers = re.findall(r'\b(\d+(?:\.\d+)?)\b', review_text)
    for num_str in standalone_numbers:
        try:
            num = float(num_str)
            if 0 <= num <= 100:
                print(f"[INFO] Using standalone number {num} as potential fit score")
                return int(round(num))
        except ValueError:
            continue
    
    print(f"[WARN] No fit score found in review text: {review_text[:200]}...")
    return 0


def get_file_content(service, file_id):
    """Helper function to get the content of a file from Google Drive."""
    try:
        # Get file metadata to check its type
        file = service.files().get(fileId=file_id, fields="mimeType, name").execute()
        mime_type = file.get('mimeType', '')
        
        # Get the file content
        data = service.files().get_media(fileId=file_id).execute()
        
        if mime_type == 'application/pdf':
            # Handle PDF files
            return extract_text(io.BytesIO(data))
            
        elif mime_type == 'application/vnd.google-apps.document':
            # For Google Docs, use the export method
            data = service.files().export(
                fileId=file_id,
                mimeType='text/plain'
            ).execute()
            return data.decode('utf-8') if isinstance(data, bytes) else str(data)
            
        else:
            # For regular text files
            if isinstance(data, bytes):
                # Try different encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        return data.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                return data.decode('utf-8', errors='ignore')
            return str(data)
            
    except Exception as e:
        raise Exception(f"Error reading file: {str(e)}")

def get_job_description_from_drive(folder_id=None):
    """Get job description from a file in Google Drive."""
    if not folder_id:
        logger.warning("No folder ID provided for job description")
        return "No folder ID provided. No job description file found in the selected folder."
    
    try:
        service = get_google_drive_service()
        
        results = service.files().list(
            q=f"'{folder_id}' in parents and (mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document') and trashed=false",
            fields="files(id, name, mimeType)",
            pageSize=1,  # We only need the first matching file
            orderBy="createdTime desc"  # Get the most recent file if multiple exist
        ).execute()

        files = results.get('files', [])
        if not files:
            logger.warning(f"No job description file found in folder {folder_id}")
            return "No job description file found in the selected folder"

        # Use the first matching file
        file_id = files[0]['id']
        file_name = files[0]['name']
        logger.info(f"Found job description file: {file_name}")
        
        return get_file_content(service, file_id)

    except Exception as e:
        error_msg = f"Error reading job description: {str(e)}"
        logger.error(error_msg)
        return error_msg

@app.get("/job_description", response_class=PlainTextResponse)
def get_job_description(folder_id: Optional[str] = None):
    """Endpoint to get the job description from Google Drive."""
    return get_job_description_from_drive(folder_id)

@app.get("/cv_content", response_class=PlainTextResponse)
def get_cv_content(source: str, path: str):
    if source == 'gdrive':
        return get_gdrive_cv_content(path)
    elif source == 'github':
        return get_github_cv_content(path)

# Promptfoo Evaluation Endpoints
evaluation_agent = CVEvaluationAgent()

class EvaluationRequest(BaseModel):
    cv_text: str
    job_requirements: str
    analysis_result: str

class EvaluationResponse(BaseModel):
    overall_score: float
    status: str
    evaluations: dict
    summary: str
    results_file: Optional[str] = None

@app.post("/evaluate/cv_analysis", response_model=EvaluationResponse)
async def evaluate_cv_analysis(request: EvaluationRequest):
    """
    Evaluate CV analysis using Promptfoo for quality, bias, and security testing.
    
    This endpoint runs comprehensive evaluation including:
    - Bias detection
    - Security validation
    - Hallucination testing
    - Output quality assessment
    """
    try:
        results = await evaluation_agent.evaluate_crew_output(
            crew_output=request.analysis_result,
            cv_text=request.cv_text,
            job_requirements=request.job_requirements,
            save_results=True
        )
        
        summary = evaluation_agent.get_evaluation_summary(results)
        
        return EvaluationResponse(
            overall_score=results.get('overall_score', 0),
            status=results.get('status', 'unknown'),
            evaluations=results.get('evaluations', {}),
            summary=summary,
            results_file=results.get('results_file')
        )
        
    except Exception as e:
        logger.error(f"Error during evaluation: {str(e)}")
        return EvaluationResponse(
            overall_score=0,
            status="error",
            evaluations={"error": str(e)},
            summary=f"Evaluation failed: {str(e)}"
        )

@app.post("/evaluate/bias_check")
async def evaluate_bias_only(analysis_text: str = Body(..., embed=True)):
    """
    Quick bias detection check for CV analysis text.
    """
    try:
        evaluator = PromptfooEvaluator()
        results = await evaluator.run_bias_detection(analysis_text)
        return results
    except Exception as e:
        return {"error": str(e), "status": "failed"}

@app.post("/evaluate/security_check")
async def evaluate_security_only(analysis_text: str = Body(..., embed=True)):
    """
    Security validation check for prompt injection and other security issues.
    """
    try:
        evaluator = PromptfooEvaluator()
        results = await evaluator.run_security_test(analysis_text)
        return results
    except Exception as e:
        return {"error": str(e), "status": "failed"}

@app.get("/evaluate/results/{filename}")
async def get_evaluation_results(filename: str):
    """
    Retrieve saved evaluation results by filename.
    """
    try:
        results_path = f"evaluations/results/{filename}"
        if os.path.exists(results_path):
            with open(results_path, 'r') as f:
                return json.load(f)
        else:
            return {"error": "Results file not found", "status": "not_found"}
    except Exception as e:
        return {"error": str(e), "status": "failed"}

@app.get("/evaluate/health")
async def evaluation_health_check():
    """
    Health check for evaluation system.
    """
    return {
        "status": "healthy",
        "promptfoo_available": True,
        "evaluation_features": [
            "bias_detection",
            "security_validation", 
            "hallucination_testing",
            "quality_assessment"
        ],
        "config_path": "evaluations/promptfooconfig.yaml"
    }
    return "Invalid source" 